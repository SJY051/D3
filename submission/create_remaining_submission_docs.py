# -*- coding: utf-8 -*-
"""Generate the six remaining D³ submission deliverables (①③④⑤⑥⑨) as editable DOCX.

Reuses the layout helpers from create_korean_submission_docs.py so all ten
deliverables share one visual system. Owner is 윤서진 for these six; 최정민 owns
②⑦⑧ (already generated).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from create_korean_submission_docs import (
    OUT,
    BLUE,
    INK,
    MUTED,
    add_body,
    add_bullets,
    add_callout,
    add_heading,
    add_table,
    configure,
    set_font,
)

OWNER = "윤서진"
REVISION = "c738cd8 (main)"
DATE = "2026. 08. 19."


def title_block(doc, title, subtitle, kind):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(kind), 10, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title), 23, True, INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run(subtitle), 12, color=MUTED)
    add_table(doc, ["프로젝트", "담당", "기준 리비전", "작성일"],
              [["D³ (Dopamin-Driven Development)", OWNER, REVISION, DATE]],
              [3100, 1700, 2100, 2100])


def _owner_footer(doc, short_title):
    # configure() writes a footer crediting 최정민; replace it with the correct owner.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run(f"D³ 제출 문서 | {OWNER} | "), 8.5, color=MUTED)
    from create_korean_submission_docs import add_page_field
    add_page_field(footer)


def source_callout(doc, rel_path, desc):
    add_callout(doc, "저장소 원본", f"{desc} — {rel_path} (https://github.com/SJY051/D3/blob/c738cd8/{rel_path})")


def requirements_doc():
    doc = Document()
    configure(doc, "요구사항 명세서")
    _owner_footer(doc, "요구사항 명세서")
    title_block(doc, "D³ 요구사항 명세서", "핵심·선택 기능 요구사항과 우선순위(MUST/SHOULD)", "제출 산출물 ① | Requirements")
    add_heading(doc, "1. 개요와 우선순위 규칙")
    add_body(doc, "D³는 실시간 코딩 배틀을 핵심(MUST, P0)으로 하고, 개발자 커뮤니티 심화 기능을 선택(SHOULD, P1)으로 둔다. P0는 시연·feature freeze 인수에 필수이고, P1은 제품 계약에 남되 P0 경로가 안정된 뒤 활성화한다. 미구현 P1 표면은 feature boundary 뒤에 남겨 mock으로 위장하지 않는다.")
    add_heading(doc, "2. 핵심 기능 요구사항 (P0 · MUST)")
    add_table(doc, ["도메인", "요구사항", "요구사항 ID"], [
        ("인증·세션", "회원가입·로그인, 세션 발급과 갱신, 계정 복구", "D3-IDN-*, D3-SEC-001"),
        ("커뮤니티", "Markdown 공개 피드 발행·조회(코드 블록 포함), 공개 전적 조회", "D3-COM-001, D3-STAT-001"),
        ("배틀", "랭크 매칭, Run/Submit, 재접속, 제출 판정 표시와 accepted 잠금, 공격 교환", "D3-BTL-001~005"),
        ("채점", "Judge0 기반 코드 실행 격리, 증거 저장, 결과 표준화", "D3-JDG-001"),
        ("점수·전적", "결과 게시물·전적·rating/RP 프로젝션", "D3-STAT-001, M-09"),
    ], [1500, 5600, 2300])
    add_heading(doc, "3. 선택 기능 요구사항 (P1 · SHOULD)")
    add_bullets(doc, [
        "팔로우·좋아요·댓글 등 커뮤니티 심화 표면 (#88, 활성화는 오늘 병합된 #112 계약부터)",
        "온보딩·튜토리얼 봇 매치 (#87)",
        "시맨틱 디자인 토큰 전면 적용과 와이어프레임 재구성 (#93)",
        "부하·카오스 테스트와 실 Judge0 애플리케이션 배포 통합",
    ])
    add_heading(doc, "4. 배점·수용 기준")
    add_table(doc, ["요소", "배점", "수용 기준"], [
        ("M-01~M-10 매트릭스", "항목별 7점", "각 요구사항 행의 observable outcome과 acceptance evidence로 판정"),
        ("P0/MUST", "시연·freeze 필수", "두 세션 골든 패스 + 수동 DB 개입 없음 + 계약·마이그레이션·UI 증거"),
        ("P1/SHOULD", "제품 계약에 유지", "feature boundary 뒤 비활성, 별도 이슈로 추적"),
    ], [1800, 1800, 5800])
    add_heading(doc, "5. 판정 증거")
    add_body(doc, "요구사항별 판정은 저장소의 권위 있는 추적 문서(docs/specs/d3-mvp.md의 내부 점수·우선순위 표, docs/artifact-status.md)와 RC c738cd8의 fresh 두-세션 인수 증거(Issue #19)를 따른다. 미실행 항목(실 Judge0 앱 통합, 부하·카오스)은 PENDING/NOT RUN으로 정직하게 표기한다.")
    source_callout(doc, "docs/specs/d3-mvp.md", "요구사항 매트릭스·배점·수용 근거의 권위 원본")
    return doc


def system_context_doc():
    doc = Document()
    configure(doc, "시스템 구성도")
    _owner_footer(doc, "시스템 구성도")
    title_block(doc, "D³ 시스템 구성도", "Discovery·Gateway·Config·Kafka 등 MSA 전체 구성", "제출 산출물 ③ | System Architecture")
    add_heading(doc, "1. 전체 토폴로지")
    add_body(doc, "로컬 P0 토폴로지는 단일 진입점인 Gateway 뒤에 4개 도메인 서비스(Identity·Battle·Judge·Community)를 두고, Discovery와 Config가 각 서비스의 발견·설정을 제공한다. 비동기 이벤트는 Kafka로, 권위 데이터는 서비스별 PostgreSQL로, 만료 가능한 조정 상태는 Redis로 관리한다.")
    add_table(doc, ["구성 요소", "역할"], [
        ("Gateway", "브라우저 단일 진입점, 인증과 라우팅, WebSocket 프록시"),
        ("Discovery", "서비스 등록·발견"),
        ("Config", "서비스별 외부 설정 공급"),
        ("Kafka", "서비스 간 버전 이벤트(match.finished, rating.changed, submission.judged 등) 전달"),
        ("PostgreSQL", "서비스별 권위 저장소, 교차 서비스 테이블 접근 금지"),
        ("Redis", "만료 가능한 매칭·연결·fan-out 조정 상태"),
        ("Judge0 host", "코드 실행 격리(운영에서는 source-SG 전용 경로)"),
    ], [2200, 7200])
    add_heading(doc, "2. 통신 경계")
    add_bullets(doc, [
        "브라우저 → Gateway만 접근, 도메인 서비스 직접 노출 없음",
        "동기: REST(커뮤니티 피드·전적, 매칭)와 WebSocket(배틀 스냅샷·명령)",
        "비동기: Kafka 이벤트(결과·rating·제출 판정)로 프로젝션 구성",
        "사용자 코드는 Judge 경계 안에서만 실행되며 소스·숨김 테스트는 비공개",
    ])
    source_callout(doc, "docs/architecture/system-context.md", "전체 토폴로지·통신 경계의 권위 원본")
    return doc


def service_doc():
    doc = Document()
    configure(doc, "서비스 구성도")
    _owner_footer(doc, "서비스 구성도")
    title_block(doc, "D³ 서비스 구성도", "마이크로서비스별 책임(Bounded Context)과 통신 방식", "제출 산출물 ④ | Service Architecture")
    add_heading(doc, "1. Bounded Context")
    add_table(doc, ["서비스", "책임", "Flyway"], [
        ("Identity", "계정·세션·인증, user-profile 이벤트 발행", "V3"),
        ("Battle", "매칭·대전 상태·제출 판정·점수/rating 확정, 이벤트 발행", "V12"),
        ("Judge", "Judge0 실행 격리·증거 저장·판정 이벤트 발행", "V3"),
        ("Community", "공개 피드·결과 게시물·전적·rating/RP·핸들 프로젝션", "V6"),
    ], [1700, 6200, 1500])
    add_heading(doc, "2. 통신 방식")
    add_table(doc, ["경로", "방식"], [
        ("Gateway → Identity", "동기 REST(가입·로그인·세션)"),
        ("Gateway → Battle", "동기 REST(대기열) + WebSocket(스냅샷·명령)"),
        ("Battle → Judge", "동기 REST(Run/Submit) + 비동기 판정 이벤트 수신"),
        ("Battle/Identity/Judge → Kafka → Community", "비동기 이벤트 프로젝션"),
    ], [3400, 6000])
    add_heading(doc, "3. 데이터·식별 경계")
    add_bullets(doc, [
        "서비스별 PostgreSQL 소유, 교차 서비스 테이블·엔티티 금지",
        "식별자·버전 계약은 교환하되 테이블은 공유하지 않음",
        "Redis는 만료 가능한 조정 상태에만 사용, PostgreSQL이 권위",
    ])
    source_callout(doc, "docs/architecture/services.md", "서비스 책임·통신·데이터 경계의 권위 원본")
    return doc


def erd_doc():
    doc = Document()
    configure(doc, "ERD")
    _owner_footer(doc, "ERD")
    title_block(doc, "D³ ERD", "서비스별 분리 DB 스키마와 관계", "제출 산출물 ⑤ | ERD")
    add_heading(doc, "1. 규모와 체인")
    add_body(doc, "논리 ERD는 서비스별로 분리된 4개 데이터베이스에 33개 테이블과 21개 서비스 내부 참조를 정의한다. 각 서비스는 forward-only Flyway 체인으로 스키마를 관리하며, 체인은 Identity V3, Community V6, Judge V3, Battle V12이다(기존 V1~이전 마이그레이션은 변경 불가).")
    add_table(doc, ["데이터베이스", "테이블 예시", "버전"], [
        ("identity", "user_account, session, refresh_session", "V3"),
        ("community", "post, match_projection, profile_projection, comment", "V6"),
        ("judge", "submission, judge_evidence, judged_outbox", "V3"),
        ("battle", "match, match_player, judge_job_reference, queue_ticket", "V12"),
    ], [2200, 5600, 1600])
    add_heading(doc, "2. 서비스 간 관계 규칙")
    add_bullets(doc, [
        "서비스 간 외래 키 없음 — 교차 데이터는 버전 이벤트로 프로젝션",
        "profile_projection·match_projection은 Community가 소유하는 파생 프로젝션",
        "judge_job_reference는 Battle이 Judge 증거를 참조하는 상관용 저장소",
    ])
    source_callout(doc, "docs/architecture/erd.dbml", "DBML 원본(렌더 가능)·테이블/참조 정의")
    return doc


def cloud_doc():
    doc = Document()
    configure(doc, "클라우드 아키텍처도")
    _owner_footer(doc, "클라우드 아키텍처도")
    title_block(doc, "D³ 클라우드 아키텍처도", "AWS 등 배포 리소스 구성", "제출 산출물 ⑥ | Cloud Architecture")
    add_heading(doc, "1. 현재 확정 범위")
    add_body(doc, "계정·리전과 zero-ingress Judge0 실행 호스트가 바인딩되어 있고, 운영용 Judge0 어댑터 6개 런타임과 source-security-group 전용 경로 스모크가 Issue #59에서 PASS했다. 임시 러너·경로는 정리됐다.")
    add_heading(doc, "2. 리소스 구성(계획)")
    add_table(doc, ["리소스", "목적", "상태"], [
        ("Judge0 host", "코드 실행 격리(고정 6개 런타임)", "바인딩·경로 스모크 PASS"),
        ("애플리케이션 서비스", "4개 도메인 서비스 배포", "PENDING (이미지·IAM·배포 미실시)"),
        ("Kafka/PostgreSQL/Redis", "이벤트·권위 데이터·조정 상태", "로컬 Compose에서 운영"),
        ("OIDC/계정", "GitHub OAuth 등", "로컬 P0에서 활성"),
    ], [2600, 4700, 2100])
    add_heading(doc, "3. 정직한 경계")
    add_bullets(doc, [
        "source-SG 전용 경로는 증명됨(#59), 배포된 judge-service가 그 경로로 실행되는 애플리케이션 통합은 NOT RUN",
        "애플리케이션 클라우드 배포·IAM·쿼터 증거는 PENDING으로 유지",
    ])
    source_callout(doc, "docs/architecture/cloud.md", "클라우드 리소스·증거 경계의 권위 원본")
    return doc


def deployment_doc():
    doc = Document()
    configure(doc, "배포 계획서")
    _owner_footer(doc, "배포 계획서")
    title_block(doc, "D³ 배포 계획서", "CI/CD 파이프라인, 배포 절차와 롤백 전략", "제출 산출물 ⑨ | Deployment Plan")
    add_heading(doc, "1. CI/CD 파이프라인")
    add_body(doc, "GitHub Actions가 PR마다 5개 게이트를 실행한다: Repository checks, Frontend build and test evidence, Backend build and test evidence, Infrastructure configuration, Delivery boundary (no deploy). 추가로 로컬 검증으로 pnpm verify:scaffold(계약·구조·마이그레이션 checksum·preflight)와 서비스별 Gradle 테스트를 돌린다.")
    add_heading(doc, "2. 배포 절차")
    add_bullets(doc, [
        "로컬: pnpm local:start(Compose)로 게이트웨이·4개 서비스·Kafka·PostgreSQL·Redis 기동, demo-preflight READY 확인",
        "제출/배포 경계: Delivery boundary 게이트가 deploy를 차단 — main은 review-only, 배포는 별도 권한",
        "Judge0: 운영 호스트·어댑터·경로는 #59에서 스모크, 애플리케이션 배포는 PENDING",
    ])
    add_heading(doc, "3. 롤백 전략")
    add_table(doc, ["상황", "전략"], [
        ("스키마", "forward-only 마이그레이션, 이전 checksum 보존으로 재생·재구축 가능"),
        ("프로젝션", "REBUILD_REQUIRED + replay queue로 이벤트 재생 재구축"),
        ("서비스", "이미지 태그 고정·롤백, Kafka auto-offset과 inbox 멱등 재처리"),
    ], [2400, 7000])
    source_callout(doc, "docs/operations/deployment-plan.md", "배포·롤백 절차의 권위 원본")
    return doc


def main():
    specs = [
        (requirements_doc, "D3_요구사항명세서.docx"),
        (system_context_doc, "D3_시스템구성도.docx"),
        (service_doc, "D3_서비스구성도.docx"),
        (erd_doc, "D3_ERD.docx"),
        (cloud_doc, "D3_클라우드아키텍처도.docx"),
        (deployment_doc, "D3_배포계획서.docx"),
    ]
    for builder, name in specs:
        builder().save(OUT / name)
        print("wrote", name)


if __name__ == "__main__":
    main()
