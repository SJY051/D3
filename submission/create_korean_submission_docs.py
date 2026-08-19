from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT = Path(__file__).parent
FONT = "Apple SD Gothic Neo"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
MUTED = "5B6573"
PAGE_WIDTH_DXA = 9360


def set_font(run, size=11, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, 9.5, True, INK)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, 10.7, True, INK)
    r = p.add_run(text)
    set_font(r, 10.7)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_font(r, 10.5)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade(cell, GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_font(r, 10.5, True, DARK_BLUE)
    r = p.add_run(text)
    set_font(r, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, 8.5, color=MUTED)


def configure(doc, short_title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(10.7)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(f"D³ | {short_title}")
    set_font(r, 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("D³ 제출 문서 | 최정민 | ")
    set_font(r, 8.5, color=MUTED)
    add_page_field(footer)


def title_block(doc, title, subtitle, kind):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(kind)
    set_font(r, 10, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, 23, True, INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_font(r, 12, color=MUTED)
    add_table(doc, ["프로젝트", "담당", "기준 리비전", "작성일"], [["D³ (Dopamin-Driven Development)", "최정민", "1afdd16 (main)", "2026. 08. 19."]], [3100, 1700, 2100, 2100])


def font(size):
    return ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", size)


def diagram(path):
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    lanes = [("사용자", "EDF4FA"), ("웹·게이트웨이", "F6F8FA"), ("도메인 서비스", "EDF4FA"), ("이벤트·커뮤니티", "F6F8FA")]
    lane_h = 190
    for i, (name, fill) in enumerate(lanes):
        y = i * lane_h
        draw.rectangle((0, y, 1800, y + lane_h), fill=f"#{fill}", outline="#D9E2EA", width=2)
        draw.text((28, y + 18), name, font=font(28), fill="#1F4D78")
    boxes = [
        (250, 58, 490, 128, "로그인·대기열 진입"),
        (650, 248, 900, 318, "인증·매칭·대전"),
        (1040, 248, 1285, 318, "Run / Submit 요청"),
        (1040, 438, 1285, 508, "Judge 증거 저장"),
        (1410, 438, 1655, 508, "Kafka 이벤트"),
        (1410, 628, 1655, 698, "결과 게시물·전적"),
    ]
    for x1, y1, x2, y2, text in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=14, fill="#FFFFFF", outline="#2E74B5", width=3)
        bbox = draw.textbbox((0, 0), text, font=font(24))
        draw.text(((x1+x2-(bbox[2]-bbox[0]))/2, (y1+y2-(bbox[3]-bbox[1]))/2-3), text, font=font(24), fill="#0B2545")
    arrows = [((490, 93), (650, 283)), ((900, 283), (1040, 283)), ((1162, 318), (1162, 438)), ((1285, 473), (1410, 473)), ((1532, 508), (1532, 628))]
    for start, end in arrows:
        draw.line((*start, *end), fill="#2E74B5", width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex-16, ey-8), (ex-16, ey+8)], fill="#2E74B5")
    draw.rounded_rectangle((235, 550, 970, 700), radius=12, fill="#FFF8E8", outline="#C99A24", width=2)
    draw.text((265, 575), "현재 종료 규칙", font=font(25), fill="#7A5A00")
    draw.text((265, 620), "양쪽 Submit만으로 종료되지 않음 → 서버 타이머 만료 후 JUDGING·점수 계산", font=font(21), fill="#5C4900")
    img.save(path)


def workflow_doc():
    doc = Document()
    configure(doc, "업무 흐름도")
    title_block(doc, "D³ 업무 흐름도", "실시간 코딩 배틀의 사용자 흐름과 MSA 서비스 책임", "제출 산출물 ② | Swimlane")
    add_heading(doc, "1. 문서 목적")
    add_body(doc, "D³의 핵심 사용자 여정은 회원 인증부터 랭크 매칭, 코드 채점, 결과·전적·커뮤니티 게시물 생성까지 이어진다. 이 문서는 각 단계의 책임 서비스와 동기/비동기 경계를 한눈에 설명한다.")
    add_callout(doc, "핵심 원칙", "브라우저는 Gateway를 통해서만 접근하고, Battle은 매치 상태와 결과를, Judge는 실행 증거를, Community는 공개 프로젝션을 각각 소유한다.")
    add_heading(doc, "2. 전체 업무 흐름", 1)
    add_body(doc, "아래 도식은 시연과 발표에서 전체 흐름을 빠르게 설명하기 위한 요약 이미지다. 세부 단계·책임·판정 기준은 바로 아래의 편집 가능한 Swimlane 표에서 수정한다.")
    image_path = OUT / "_workflow_overview.png"
    diagram(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.35))
    image_path.unlink(missing_ok=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("그림 1. D³ 랭크 배틀 전체 업무 흐름")
    set_font(r, 9, color=MUTED)
    add_body(doc, "아래 Swimlane 표는 Word 표로 작성되어 각 서비스명·단계·설명 문구를 직접 수정할 수 있다.")
    add_table(doc, ["단계", "사용자", "웹·Gateway", "도메인·이벤트 서비스"], [
        ("1", "로그인·대기열 진입", "인증 요청과 보호 API 전달", "Identity 세션 발급, Battle 대기열 등록"),
        ("2", "Ready·코드 작성", "매치 상태와 서버 시계 표시", "Battle이 매치·연결·공격 상태 관리"),
        ("3", "Run·Submit", "Judge 요청·판정 표시", "Judge 실행 증거 저장 → Battle에 참조 전달"),
        ("4", "결과·전적 조회", "결과 화면·Feed로 이동", "마감 후 점수 확정 → Kafka → Community 프로젝션"),
    ], [650, 1900, 2500, 4310])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("표 1. D³ 랭크 배틀 업무 흐름(Swimlane)")
    set_font(r, 9, color=MUTED)
    add_heading(doc, "3. 단계별 Swimlane", 1)
    add_table(doc, ["단계", "사용자·웹", "서비스 책임", "확인 결과"], [
        ("1. 인증", "회원가입 또는 로그인", "Identity가 세션을 발급하고 Gateway가 보호된 API를 전달", "로그인 후 Feed 접근"),
        ("2. 매칭", "언어 선택 후 Ranked Queue 진입", "Battle이 동일 언어·근접 rating을 기준으로 매칭", "두 세션이 같은 matchId 수신"),
        ("3. 대전", "Ready 후 코드 작성, Run·Submit", "Battle이 서버 시계·연결 상태·공격을 권위 있게 관리", "상대 소스는 마스킹"),
        ("4. 채점", "공개 예제 Run, 숨김 테스트 Submit", "Judge가 상태·통과 수·실행시간 증거를 저장", "로컬은 deterministic fake Judge"),
        ("5. 종료·점수", "결과 화면 확인", "Battle이 마감 시각 뒤 증거로 승패·rating·RP를 확정", "정확한 점수 구성 저장"),
        ("6. 프로젝션", "Feed·전적 조회", "Kafka 이벤트를 Community가 소비하여 결과 게시물·전적 생성", "원본 코드·숨김 테스트 미노출"),
    ], [800, 2380, 3550, 2630])
    add_heading(doc, "4. 현재 매치 종료 규칙", 1)
    add_body(doc, "현재 구현에서는 한쪽 또는 양쪽 플레이어가 Submit을 눌러도 매치가 즉시 끝나지 않는다. Submit은 Judge 증거를 기록하며, 서버가 가진 매치 타이머가 만료되면 Battle이 RUNNING에서 JUDGING으로 전환한 뒤 저장된 증거를 비교하여 FINISHED를 확정한다.")
    add_bullets(doc, [
        "한 명만 정답이면 정답 플레이어가 승리한다.",
        "둘 다 정답이면 풀이 속도(50%), 동적 실행 효율(35%), 제출 절제(15%)를 비교한다.",
        "둘 다 정답이 아니면 hidden test 통과율을 주 점수로 사용한다.",
        "항복, 재접속 제한시간 초과, 플랫폼 장애는 별도 종료 경로이며 플랫폼 장애는 무효 처리한다.",
    ])
    add_heading(doc, "5. 신뢰·개인정보 경계", 1)
    add_bullets(doc, [
        "사용자 코드는 Judge 경계 밖으로 공개되지 않으며, 상대방에게는 연결·진행 상태만 제한적으로 전달한다.",
        "PostgreSQL은 결과·점수·전적의 권위 저장소이며 Redis는 만료 가능한 매칭·연결·fan-out 상태에만 사용한다.",
        "Community는 Battle 데이터베이스를 조회하지 않고 버전이 있는 Kafka 이벤트를 자체 프로젝션으로 소비한다.",
    ])
    add_heading(doc, "6. 핵심 MSA 요구사항 대응", 1)
    add_table(doc, ["핵심 요건", "D³ 반영 내용"], [
        ("서비스 분리", "Identity·Battle·Judge·Community가 각 데이터와 책임을 분리한다."),
        ("인프라 경계", "Gateway·Discovery·Config를 통해 단일 진입점과 서비스 설정 관리를 제공한다."),
        ("비동기·데이터", "Kafka 이벤트와 서비스별 PostgreSQL 소유로 교차 DB 조회를 금지한다."),
    ], [2500, 6860])
    return doc


def wireframe_doc():
    doc = Document()
    configure(doc, "화면 기획서")
    title_block(doc, "D³ 화면 기획서", "P0 Golden Path 화면과 P1 기능 경계", "제출 산출물 ⑦ | Wireframe·기능 정의")
    add_heading(doc, "1. 화면 설계 방향")
    add_body(doc, "D³는 다크 퍼스트, 정보 밀도 중심의 화면을 사용한다. 커뮤니티·전적 화면은 360px에서도 읽기 쉽게 구성하고, 코드 편집을 포함하는 배틀 화면은 1280px 데스크톱을 기준으로 한다.")
    add_table(doc, ["영역", "사용자 흐름", "현재 상태"], [
        ("공개 영역", "로그인 → Feed → 결과/전적", "P0 구현"),
        ("랭크 영역", "언어 선택 → 대기열 → 배틀 → 결과 → 전적", "P0 구현"),
        ("연습 영역", "문제 목록 → 편집기 → 개인 풀이", "P1 구조 프로토타입"),
        ("운영 영역", "문제 목록 → 검토 → 활성화", "P1 구조 프로토타입"),
    ], [1650, 4800, 2910])
    add_heading(doc, "2. P0 화면별 기능 정의", 1)
    add_table(doc, ["ID", "화면·경로", "핵심 기능", "사용자 확인 사항"], [
        ("WF-01", "로그인 /sign-in", "이메일·비밀번호 로그인, 계정 생성", "인증 오류와 입력 규칙을 텍스트로 안내"),
        ("WF-02", "커뮤니티 Feed /feed", "공개 Markdown 게시, fenced code 렌더링, 결과 게시물 확인", "코드 블록은 본문 글자 수와 분리, 공개 범위만 노출"),
        ("WF-03", "랭크 대기열 /ranked", "언어 선택, QUEUED→MATCHED 상태, 배틀 이동", "대기 상태와 경과 시간을 색상 외 텍스트로 표시"),
        ("WF-04", "실시간 배틀 /battles/:matchId", "문제·에디터·Run/Submit·공격·상대 마스킹", "서버 타이머·연결 상태·Judge 상태를 명시"),
        ("WF-05", "결과 /results/:matchId", "승패·무승부·무효, 결과 전적, Feed/전적 이동", "viewer-relative 결과와 공개 match record"),
        ("WF-06", "플레이어 전적 /players/:playerId", "최근 ACTIVE 경기, keyset 더보기", "공개 전적만 노출, 개인 소스 미노출"),
    ], [720, 2050, 3550, 3040])
    add_heading(doc, "3. WF-04 배틀 화면의 상태 설계", 1)
    add_bullets(doc, [
        "문제 영역: 문제 설명, 제약, 공개 예제와 진단 결과를 표시한다.",
        "에디터 영역: 사용자의 소스만 수정 가능하며 공격 오버레이는 저장된 소스를 바꾸지 않는다.",
        "상대 정보: 연결 상태·제한된 진행 상태만 마스킹하여 표시하고, 상대 코드·리터럴·식별자는 노출하지 않는다.",
        "행동 영역: Ready, Run, Submit, 공격·방어·반사 상태를 제공한다. Submit은 현재 매치를 종료하지 않으며 Judge 증거를 기록한다.",
        "종료: 서버 타이머가 만료된 후 결과 계산 단계로 이동한다. 재접속, 항복, 플랫폼 장애는 별도 상태로 안내한다.",
    ])
    add_heading(doc, "4. 공통 셸과 접근성", 1)
    add_body(doc, "로그인한 사용자가 진행 중인 매치를 보유하면, 배틀 화면 이외의 P0 화면 상단에 ‘진행 중인 매치로 돌아가기’ 배너를 표시한다. 매치가 종료되거나 소유자가 변경되면 배너는 사라진다.")
    add_bullets(doc, [
        "키보드만으로 주요 입력·버튼에 접근할 수 있고 focus 표시를 유지한다.",
        "타이머, 연결, Judge 결과, 공격 경고는 색상만으로 전달하지 않고 텍스트·아이콘을 함께 사용한다.",
        "reduced-motion 환경에서는 대기열의 움직임을 줄이거나 정적인 상태 표시로 대체한다.",
    ])
    add_heading(doc, "5. P1 기능 경계", 1)
    add_callout(doc, "제출 시 유의", "WF-07 문제 운영과 WF-08 Solo Practice는 현재 라우트·구조만 제공하는 P1 프로토타입이다. 실제 문제 관리, 연습 채점, 개인 풀이 보관 기능이 구현된 것처럼 시연하거나 설명하지 않는다.")
    add_table(doc, ["ID", "화면", "P1 경계"], [
        ("WF-07", "/admin/problems", "문제 목록·검토·발행 영역의 구조 화면이며 운영 워크플로는 활성화되지 않음"),
        ("WF-08", "/practice", "연습 편집기 구조 화면이며 Submit 실행·개인 풀이 저장은 활성화되지 않음"),
    ], [950, 2100, 6570])
    add_heading(doc, "6. 기능 우선순위(MUST / SHOULD)", 1)
    add_table(doc, ["우선순위", "기능 범위", "제출 시 판정"], [
        ("MUST", "로그인, Feed, 랭크 매칭, 배틀, 결과·전적", "P0 구현·시연 대상"),
        ("SHOULD", "문제 운영, Solo Practice, 확장 커뮤니티", "P1 구조만 제공하며 미구현으로 명시"),
    ], [1500, 4700, 3160])
    return doc


def testplan_doc():
    doc = Document()
    configure(doc, "테스트 계획서")
    title_block(doc, "D³ 테스트 계획서", "단위·통합·브라우저·부하·카오스 검증 계획", "제출 산출물 ⑧ | QA 계획·증거 기준")
    add_heading(doc, "1. 테스트 목적과 원칙")
    add_body(doc, "D³는 인증, 실시간 대전, Judge 경계, 결과 프로젝션이 분리된 MSA 구조이므로 각 서비스 단위의 규칙 검증과 서비스 간 계약 검증을 함께 수행한다. 모든 결과는 PASS, FAIL, SKIP, NOT RUN으로 분리 기록하며, 로컬 fake Judge 증거와 실제 Judge0 증거를 혼동하지 않는다.")
    add_callout(doc, "증거 원칙", "기능이 화면에 보인다는 사실만으로 완료 처리하지 않는다. 실행 명령, 리비전, 환경, 관찰 결과, 오류·상관관계 식별자를 함께 남긴다.")
    add_heading(doc, "2. 테스트 레이어", 1)
    add_table(doc, ["레이어", "주요 검증 대상", "현재 증거", "상태"], [
        ("도메인 단위", "매치 상태·점수·rating·공격", "결정론적 시계와 고정 경계 테스트", "PARTIAL PASS"),
        ("어댑터 통합", "PostgreSQL·Redis·Kafka·outbox/inbox", "컨테이너 기반 migration·중복·재시도 검증", "PARTIAL PASS"),
        ("계약", "HTTP·WebSocket·이벤트", "버전·권한·개인정보 음성 케이스", "PARTIAL PASS"),
        ("브라우저", "랭크 Golden Path·소스 보존", "WF-04 credential·공격 오버레이·8/18 리허설", "PARTIAL PASS"),
        ("Judge0 호스트", "6개 언어 runtime·격리", "지정 호스트 adapter/host smoke", "PASS 경계"),
        ("부하·카오스", "fan-out·queue·복구", "계획 수립", "NOT RUN"),
    ], [1400, 2750, 3500, 1710])
    add_heading(doc, "3. 핵심 기능 시나리오", 1)
    add_table(doc, ["시나리오", "검증 흐름", "판정 기준", "현재 상태"], [
        ("A. Golden Path", "두 계정 로그인 → 같은 언어 매칭 → Ready → Run/Submit → 결과·전적·게시물", "DB 수동 수정 없이 두 세션이 동일 매치와 결과를 확인", "V12 fresh capture PENDING"),
        ("B. 재접속·항복", "RUNNING 중 연결 종료 후 30초 내/외 복귀, 항복", "서버 시계 유지, 복귀 또는 정확히 한 번의 종료", "재접속 일부 증거, 항복 UI NOT RUN"),
        ("C. 플랫폼 장애", "Judge platform failure", "매치 VOIDED, rating·RP 미변경", "도메인/계약 PASS, 라이브 NOT RUN"),
        ("D. 커뮤니티·개인정보", "Markdown fenced code 게시, 결과 record 조회", "소스·숨김 테스트가 공개 API에 없음", "서비스 경계 PASS, 최종 화면 캡처 PENDING"),
    ], [1370, 3630, 2850, 1510])
    add_heading(doc, "4. Judge 및 점수 검증 기준", 1)
    add_bullets(doc, [
        "Judge 상태는 ACCEPTED, WRONG_ANSWER, COMPILATION_ERROR, RUNTIME_ERROR, TIME_LIMIT, MEMORY_LIMIT, PLATFORM_FAILURE로 구분한다.",
        "로컬 deterministic fake Judge는 정해진 입력으로 결과를 재현한다. 실제 임의 코드를 실행한 Judge0 결과로 설명하지 않는다.",
        "한 명만 정답이면 즉시 승패 우선 규칙을 적용한다. 둘 다 정답이면 속도 50%, 동적 실행 효율 35%, 제출 절제 15%를 비교한다.",
        "둘 다 미정답이면 hidden test 통과율을 주 점수로 사용한다. 정확히 같으면 무승부다.",
        "현재 Submit은 증거 기록이며 종료 명령이 아니다. 서버 타이머 종료 후 JUDGING·결과 계산이 수행되는지 확인한다.",
    ])
    add_heading(doc, "5. 시연 전 Preflight", 1)
    add_table(doc, ["순서", "점검", "통과 기준"], [
        ("1", "의존성·구조 검증", "pnpm install --frozen-lockfile 및 pnpm verify:scaffold 성공"),
        ("2", "로컬 기동", "pnpm local:start 후 demo-preflight: READY"),
        ("3", "서비스 상태", "Web, Gateway, Identity, Battle, Judge, Community, PostgreSQL, Redis, Kafka 확인"),
        ("4", "데모 문제", "V12가 demo-sum-v1을 fresh DB에 자동 seed; 수동 INSERT 금지"),
        ("5", "시연 세션", "독립 브라우저 2개와 각 계정의 로그인 상태 준비"),
    ], [700, 2600, 6060])
    add_heading(doc, "6. 부하·카오스 계획", 1)
    add_table(doc, ["구분", "주입/부하", "관찰값", "현재 상태"], [
        ("WebSocket fan-out", "paired room 단계적 증가", "연결 성공률, p95/p99, Redis ops", "NOT RUN"),
        ("Judge queue", "6개 언어 staged submit", "대기시간, 실행시간, 실패 분류", "NOT RUN"),
        ("Redis restart", "대기열·presence 중 Redis 재시작", "PostgreSQL 결과 보존·재수렴", "NOT RUN"),
        ("Kafka replay", "이벤트 pause/duplicate 뒤 재개", "inbox idempotency·projection 단일성", "NOT RUN"),
        ("Judge outage", "platform failure/연결 종료", "VOIDED·rating/RP 변화 없음", "NOT RUN"),
    ], [1600, 2800, 3000, 1960])
    add_heading(doc, "7. 최종 기록 양식", 1)
    add_bullets(doc, [
        "리비전 SHA, 실행 시각·시간대, 운영자, local/deployed 환경, 선택 Judge adapter를 기록한다.",
        "각 항목마다 PASS/FAIL/SKIP/NOT RUN을 별도 집계하고, 실패 시 요청 또는 correlation ID와 민감정보가 제거된 로그를 보관한다.",
        "최종 영상에는 preflight READY, 실제 서비스 화면, 결과·전적·자동 게시물을 순서대로 담고, fake Judge 또는 fallback 여부를 화면과 구두로 명확히 표시한다.",
    ])
    add_heading(doc, "8. 공통 요구사항 검증 대응", 1)
    add_table(doc, ["구분", "검증 대상", "기록 방식"], [
        ("MUST", "서비스 분리·Gateway·Discovery·Config·Kafka·서비스별 DB", "계약·통합 테스트 결과와 증거 링크"),
        ("SHOULD", "모니터링·Circuit Breaker·중앙 로그 등 선택 항목", "선택 여부와 PASS 또는 NOT RUN을 명시"),
    ], [1300, 4800, 3260])
    return doc


def kfont(size, bold=False):
    return ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", size, index=8 if bold else 0)


def wrap(draw, text, font_obj, width):
    words, lines, line = text.split(" "), [], ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if draw.textlength(candidate, font=font_obj) <= width:
            line = candidate
            continue
        if line:
            lines.append(line)
            line = word
        while draw.textlength(line, font=font_obj) > width:
            cut = len(line) - 1
            while cut > 1 and draw.textlength(line[:cut], font=font_obj) > width:
                cut -= 1
            lines.append(line[:cut])
            line = line[cut:]
    if line:
        lines.append(line)
    return lines


def page_canvas(title, subtitle, page):
    img = Image.new("RGB", (1700, 2200), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1700, 16), fill="#2E74B5")
    draw.text((130, 78), "D³ 제출 문서 | 최정민", font=kfont(24), fill="#5B6573")
    draw.text((130, 145), title, font=kfont(60, True), fill="#0B2545")
    draw.text((130, 225), subtitle, font=kfont(30), fill="#5B6573")
    draw.line((130, 290, 1570, 290), fill="#D7E2EF", width=3)
    draw.text((130, 2115), "D³ (Dopamin-Driven Development) | 2026. 08. 19.", font=kfont(21), fill="#687385")
    draw.text((1570, 2115), str(page), font=kfont(21), fill="#687385", anchor="ra")
    return img, draw, 345


def page_heading(draw, y, text):
    draw.text((130, y), text, font=kfont(38, True), fill="#2E74B5")
    return y + 64


def page_paragraph(draw, y, text, width=1440, size=26, color="#1E2935", leading=13):
    font_obj = kfont(size)
    for line in wrap(draw, text, font_obj, width):
        draw.text((130, y), line, font=font_obj, fill=color)
        y += size + leading
    return y + 20


def page_bullets(draw, y, items, size=25):
    for item in items:
        draw.ellipse((138, y + 14, 150, y + 26), fill="#2E74B5")
        lines = wrap(draw, item, kfont(size), 1370)
        for line in lines:
            draw.text((175, y), line, font=kfont(size), fill="#1E2935")
            y += size + 11
        y += 8
    return y


def page_table(draw, y, headers, rows, widths, size=21):
    x0, total = 130, sum(widths)
    head_h = 54
    x = x0
    for header, w in zip(headers, widths):
        draw.rectangle((x, y, x + w, y + head_h), fill="#E8EEF5", outline="#7A8FA5", width=2)
        draw.text((x + 13, y + 13), header, font=kfont(size, True), fill="#16385F")
        x += w
    y += head_h
    for row in rows:
        blocks = [wrap(draw, str(value), kfont(size), w - 26) for value, w in zip(row, widths)]
        row_h = max(52, max(len(block) for block in blocks) * (size + 9) + 22)
        x = x0
        for block, w in zip(blocks, widths):
            draw.rectangle((x, y, x + w, y + row_h), fill="white", outline="#8D9DAF", width=2)
            ty = y + 11
            for line in block:
                draw.text((x + 13, ty), line, font=kfont(size), fill="#1E2935")
                ty += size + 9
            x += w
        y += row_h
    return y + 28


def save_raster_docx(name, pages):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.header_distance = Inches(0.1)
    section.footer_distance = Inches(0.1)
    files = []
    for index, image in enumerate(pages, 1):
        path = OUT / f"_{name}_{index}.png"
        image.save(path, quality=95)
        files.append(path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.add_run().add_picture(str(path), width=Inches(7.65))
        if index < len(pages):
            doc.add_page_break()
    output = OUT / f"{name}.docx"
    doc.core_properties.author = "최정민"
    doc.core_properties.title = name
    doc.core_properties.subject = "D³ 제출 산출물"
    doc.save(output)
    for path in files:
        path.unlink(missing_ok=True)


def raster_workflow_pages():
    p1, d, y = page_canvas("D³ 업무 흐름도", "실시간 코딩 배틀의 사용자 흐름과 MSA 서비스 책임", 1)
    y = page_heading(d, y, "1. 문서 목적")
    y = page_paragraph(d, y, "D³는 회원 인증부터 랭크 매칭, 코드 채점, 결과·전적·커뮤니티 게시물 생성까지 하나의 흐름으로 연결하는 실시간 코딩 배틀 플랫폼이다. 이 문서는 단계별 서비스 책임과 데이터 흐름을 설명한다.")
    d.rounded_rectangle((130, y, 1570, y + 105), radius=14, fill="#F2F6FA", outline="#B7C9DD", width=2)
    d.text((160, y + 22), "핵심 원칙", font=kfont(26, True), fill="#1F4D78")
    d.text((330, y + 22), "Gateway는 단일 진입점이며, Battle·Judge·Community는 각자 데이터를 소유한다.", font=kfont(25), fill="#1E2935")
    y += 145
    y = page_heading(d, y, "2. 전체 업무 흐름")
    path = OUT / "_diagram.png"; diagram(path)
    flow = Image.open(path).resize((1440, 608))
    p1.paste(flow, (130, y)); path.unlink(missing_ok=True)
    return [p1, raster_workflow_page2()]


def raster_workflow_page2():
    img, d, y = page_canvas("D³ 업무 흐름도", "단계별 Swimlane 및 현재 종료 규칙", 2)
    y = page_heading(d, y, "3. 단계별 Swimlane")
    y = page_table(d, y, ["단계", "사용자·웹", "서비스 책임", "확인 결과"], [
        ("1. 인증", "회원가입·로그인", "Identity 세션 발급, Gateway 전달", "Feed 접근"),
        ("2. 매칭", "언어 선택·대기열", "Battle이 언어·rating 기준 매칭", "같은 matchId"),
        ("3. 대전", "Ready, 코드 작성, Run·Submit", "Battle이 서버 시계·연결·공격 관리", "상대 소스 마스킹"),
        ("4. 채점", "Run·Submit", "Judge가 상태·통과 수·실행 증거 저장", "로컬 fake Judge"),
        ("5. 결과", "결과·전적 확인", "Battle이 승패·rating·RP 확정", "점수 구성 저장"),
        ("6. 프로젝션", "Feed·전적 조회", "Kafka 이벤트를 Community가 소비", "결과 게시물 생성"),
    ], [220, 390, 500, 330], size=19)
    y = page_heading(d, y, "4. 현재 매치 종료 규칙")
    y = page_paragraph(d, y, "현재 Submit은 Judge 증거를 기록하는 동작이다. 한쪽 또는 양쪽이 Submit을 눌러도 매치는 즉시 종료되지 않으며, 서버 타이머가 만료된 뒤 RUNNING → JUDGING → FINISHED 순서로 결과를 확정한다.", size=24)
    y = page_bullets(d, y, ["한 명만 정답이면 정답 플레이어가 승리한다.", "둘 다 정답이면 풀이 속도 50%, 동적 실행 효율 35%, 제출 절제 15%를 비교한다.", "플랫폼 장애는 VOIDED로 처리하며 rating·RP를 변경하지 않는다."], size=23)
    y = page_heading(d, y, "5. 신뢰·개인정보 경계")
    y = page_bullets(d, y, ["PostgreSQL은 결과·점수의 권위 저장소이며 Redis는 만료 가능한 상태에만 사용한다.", "Community는 다른 서비스 DB를 조회하지 않고 버전이 있는 Kafka 이벤트로 자체 프로젝션을 만든다.", "사용자 코드·숨김 테스트·자격 증명은 공개 이벤트와 전적에 포함하지 않는다."], size=22)
    y = page_heading(d, y, "6. 핵심 MSA 요구사항 대응")
    page_table(d, y, ["핵심 요건", "D³ 반영 내용"], [
        ("서비스 분리", "Identity·Battle·Judge·Community가 각 데이터와 책임을 분리"),
        ("인프라 경계", "Gateway·Discovery·Config를 통한 단일 진입점과 서비스 설정 관리"),
        ("비동기·데이터", "Kafka 이벤트와 서비스별 PostgreSQL 소유로 교차 DB 조회 금지"),
    ], [410, 1030], size=17)
    return img


def raster_wireframe_pages():
    img, d, y = page_canvas("D³ 화면 기획서", "P0 Golden Path 화면과 P1 기능 경계", 1)
    y = page_heading(d, y, "1. 화면 설계 방향")
    y = page_paragraph(d, y, "D³는 다크 퍼스트와 정보 밀도를 기본으로 한다. 커뮤니티·전적 화면은 360px에서도 읽기 쉽게 구성하고, 코드 편집을 포함한 배틀 화면은 1280px 데스크톱을 기준으로 한다.")
    y = page_table(d, y, ["영역", "사용자 흐름", "상태"], [("공개", "로그인 → Feed → 결과/전적", "P0 구현"), ("랭크", "언어 선택 → 대기열 → 배틀 → 결과", "P0 구현"), ("연습", "문제 목록 → 편집기 → 개인 풀이", "P1 구조"), ("운영", "문제 목록 → 검토 → 활성화", "P1 구조")], [260, 720, 460], size=20)
    y = page_heading(d, y, "2. P0 화면별 기능 정의")
    y = page_table(d, y, ["ID", "화면·경로", "핵심 기능"], [("WF-01", "로그인 /sign-in", "로그인·계정 생성, 인증 오류 안내"), ("WF-02", "Feed /feed", "공개 Markdown 게시, fenced code, 결과 게시물"), ("WF-03", "랭크 /ranked", "언어 선택, QUEUED → MATCHED 상태, 배틀 이동")], [210, 430, 800], size=20)
    return [img, raster_wireframe_page2()]


def raster_wireframe_page2():
    img, d, y = page_canvas("D³ 화면 기획서", "배틀 상태 설계·접근성·P1 경계", 2)
    y = page_heading(d, y, "2. P0 화면별 기능 정의 (계속)")
    y = page_table(d, y, ["ID", "화면·경로", "핵심 기능"], [("WF-04", "배틀 /battles/:matchId", "문제·에디터·Run/Submit·공격·상대 마스킹"), ("WF-05", "결과 /results/:matchId", "승패·무승부·무효, 결과 전적, Feed 이동"), ("WF-06", "전적 /players/:playerId", "최근 ACTIVE 경기, keyset 더보기")], [210, 430, 800], size=20)
    y = page_heading(d, y, "3. WF-04 배틀 화면의 상태 설계")
    y = page_bullets(d, y, ["문제·에디터·테스트 진단과 Run/Submit 행동 영역을 분리한다.", "상대 정보는 연결·제한된 진행 상태만 표시하며 코드·리터럴·식별자는 노출하지 않는다.", "Submit은 현재 매치를 종료하지 않고 Judge 증거를 기록한다. 서버 타이머 만료 후 결과 계산 단계로 전환한다.", "공격 오버레이는 저장된 소스를 변경하지 않으며 warning·block·reflect 상태를 텍스트와 함께 표시한다."], size=23)
    y = page_heading(d, y, "4. 공통 셸과 접근성")
    y = page_paragraph(d, y, "진행 중인 매치를 가진 사용자는 배틀 화면 외 P0 화면에서 ‘진행 중인 매치로 돌아가기’ 배너를 확인할 수 있다. 주요 조작은 키보드 접근, focus 표시, 색상 외 상태 텍스트, reduced-motion 대체 표시를 제공한다.", size=23)
    y = page_heading(d, y, "5. P1 기능 경계")
    y = page_table(d, y, ["ID", "화면", "제출 시 설명 기준"], [("WF-07", "/admin/problems", "문제 관리 구조 화면이며 실제 운영 워크플로는 미활성"), ("WF-08", "/practice", "연습 편집기 구조 화면이며 Submit 실행·풀이 보관은 미활성")], [220, 420, 800], size=20)
    y = page_heading(d, y, "6. 기능 우선순위(MUST / SHOULD)")
    page_table(d, y, ["우선순위", "기능 범위", "제출 시 판정"], [
        ("MUST", "로그인, Feed, 랭크 매칭, 배틀, 결과·전적", "P0 구현·시연 대상"),
        ("SHOULD", "문제 운영, Solo Practice, 확장 커뮤니티", "P1 구조만 제공; 미구현으로 명시"),
    ], [250, 670, 750], size=18)
    return img


def raster_test_pages():
    img, d, y = page_canvas("D³ 테스트 계획서", "단위·통합·브라우저·부하·카오스 검증 계획", 1)
    y = page_heading(d, y, "1. 테스트 목적과 원칙")
    y = page_paragraph(d, y, "D³는 인증, 실시간 대전, Judge 경계, 결과 프로젝션이 분리된 MSA 구조다. 기능 화면만으로 완료를 판단하지 않고 실행 명령, 리비전, 환경, 관찰 결과를 함께 기록한다.")
    d.rounded_rectangle((130, y, 1570, y + 90), radius=14, fill="#F2F6FA", outline="#B7C9DD", width=2)
    d.text((160, y + 21), "증거 원칙", font=kfont(25, True), fill="#1F4D78")
    d.text((330, y + 22), "PASS, FAIL, SKIP, NOT RUN을 분리하고 fake Judge와 실제 Judge0 증거를 혼동하지 않는다.", font=kfont(23), fill="#1E2935")
    y += 125
    y = page_heading(d, y, "2. 테스트 레이어")
    y = page_table(d, y, ["레이어", "주요 검증", "상태"], [("도메인", "매치 상태·점수·rating·공격", "PARTIAL PASS"), ("어댑터", "PostgreSQL·Redis·Kafka·outbox/inbox", "PARTIAL PASS"), ("계약", "HTTP·WebSocket·이벤트·권한", "PARTIAL PASS"), ("브라우저", "Golden Path·소스 보존", "PARTIAL PASS"), ("Judge0 호스트", "6개 runtime·격리", "PASS 경계"), ("부하·카오스", "fan-out·queue·복구", "NOT RUN")], [330, 770, 340], size=19)
    y = page_heading(d, y, "3. 핵심 기능 시나리오")
    page_table(d, y, ["시나리오", "검증 흐름", "현재 상태"], [("A. Golden Path", "두 계정 로그인 → 매칭 → Run/Submit → 결과·전적·게시물", "V12 fresh capture PENDING"), ("B. 재접속·항복", "30초 내/외 복귀, 항복", "재접속 일부 증거"), ("C. 플랫폼 장애", "Judge platform failure → VOIDED", "라이브 NOT RUN"), ("D. 커뮤니티·개인정보", "Markdown, fenced code, match record", "최종 화면 캡처 PENDING")], [330, 720, 390], size=18)
    return [img, raster_test_page2()]


def raster_test_page2():
    img, d, y = page_canvas("D³ 테스트 계획서", "Judge·점수 검증과 시연 전 품질 게이트", 2)
    y = page_heading(d, y, "4. Judge 및 점수 검증 기준")
    y = page_bullets(d, y, ["Judge 상태는 ACCEPTED, WRONG_ANSWER, COMPILATION_ERROR, RUNTIME_ERROR, TIME_LIMIT, MEMORY_LIMIT, PLATFORM_FAILURE로 구분한다.", "로컬 deterministic fake Judge는 정해진 입력으로 결과를 재현한다. 실제 임의 코드 실행 결과로 설명하지 않는다.", "한 명만 정답이면 승리한다. 둘 다 정답이면 속도 50%, 동적 실행 효율 35%, 제출 절제 15%를 비교한다.", "현재 Submit은 증거 기록이다. 서버 타이머가 만료된 후 JUDGING·결과 계산이 진행되는지 확인한다."], size=22)
    y = page_heading(d, y, "5. 시연 전 Preflight")
    y = page_table(d, y, ["순서", "점검", "통과 기준"], [("1", "의존성·구조", "pnpm install --frozen-lockfile, pnpm verify:scaffold 성공"), ("2", "로컬 기동", "pnpm local:start 후 demo-preflight: READY"), ("3", "서비스 상태", "Web·Gateway·Identity·Battle·Judge·Community·DB·Redis·Kafka"), ("4", "데모 문제", "V12 demo-sum-v1 자동 seed, 수동 INSERT 금지"), ("5", "시연 세션", "독립 브라우저 2개와 로그인 상태 준비")], [120, 360, 960], size=18)
    y = page_heading(d, y, "6. 부하·카오스 계획")
    y = page_table(d, y, ["구분", "관찰값", "상태"], [("WebSocket fan-out", "연결 성공률, p95/p99, Redis ops", "NOT RUN"), ("Judge queue", "대기시간, 실행시간, 실패 분류", "NOT RUN"), ("Redis restart", "PostgreSQL 결과 보존·재수렴", "NOT RUN"), ("Kafka replay", "inbox idempotency·projection 단일성", "NOT RUN"), ("Judge outage", "VOIDED·rating/RP 미변경", "NOT RUN")], [400, 700, 340], size=18)
    y = page_heading(d, y, "7. 최종 기록")
    y = page_bullets(d, y, ["리비전 SHA, 시각·시간대, 환경, 선택 Judge adapter, PASS/FAIL/SKIP/NOT RUN을 기록한다.", "최종 영상에는 preflight READY와 실제 서비스 화면을 담고 fake Judge·fallback 여부를 명확히 표시한다."], size=21)
    y = page_heading(d, y, "8. 공통 요구사항 검증 대응")
    page_table(d, y, ["구분", "검증 대상", "기록 방식"], [
        ("MUST", "서비스 분리·Gateway·Discovery·Config·Kafka·서비스별 DB", "계약·통합 테스트 결과와 증거 링크"),
        ("SHOULD", "모니터링·Circuit Breaker·중앙 로그 등 선택 항목", "선택 여부와 PASS 또는 NOT RUN을 명시"),
    ], [210, 650, 810], size=17)
    return img


def main():
    workflow_doc().save(OUT / "D3_업무흐름도.docx")
    wireframe_doc().save(OUT / "D3_화면기획서.docx")
    testplan_doc().save(OUT / "D3_테스트계획서.docx")


if __name__ == "__main__":
    main()
